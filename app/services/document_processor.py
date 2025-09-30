import fitz  # PyMuPDF (for PDFs)
from openai import OpenAI
import json
import os
import base64
import time
from docx import Document
from typing import Optional, Dict, Any
from sqlalchemy.orm import Session
from app.models import DynamicPrompt, ProcessedDocument
from app.logger import get_logger

logger = get_logger(__name__)

class DocumentProcessor:
    def __init__(self, openai_api_key: str):
        self.client = OpenAI(api_key=openai_api_key)
    
    def extract_text_from_pdf(self, pdf_path: str, model: str = "gpt-4o-mini") -> str:
        """Extract text from both text-based and scanned PDFs using OpenAI Vision when needed."""
        try:
            doc = fitz.open(pdf_path)
            full_text = ""
            total_pages = len(doc)
            
            logger.info(f"Processing PDF with {total_pages} pages")

            for page_num, page in enumerate(doc, start=1):
                try:
                    text = page.get_text("text")

                    if text.strip():  
                        logger.info(f"[Page {page_num}/{total_pages}] Text detected directly")
                        full_text += text + "\n"
                    else:
                        logger.info(f"[Page {page_num}/{total_pages}] No text found, using OpenAI Vision OCR...")
                        pix = page.get_pixmap(dpi=200)  # render page as image
                        img_bytes = pix.tobytes("png")
                        img_b64 = base64.b64encode(img_bytes).decode("utf-8")

                        response = self.client.chat.completions.create(
                            model=model,
                            messages=[
                                {
                                    "role": "user",
                                    "content": [
                                        {"type": "text", "text": "Extract all text clearly from this image."},
                                        {
                                            "type": "image_url",
                                            "image_url": {"url": f"data:image/png;base64,{img_b64}"},
                                        },
                                    ],
                                }
                            ],
                        )

                        page_text = response.choices[0].message.content.strip()
                        full_text += page_text + "\n"
                        logger.info(f"[Page {page_num}/{total_pages}] OCR processing completed")
                        
                except Exception as e:
                    logger.error(f"Error processing page {page_num}: {str(e)}")
                    # Continue with next page instead of failing completely
                    continue
            
            doc.close()
            logger.info(f"PDF processing completed. Extracted {len(full_text)} characters")
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"Error opening PDF file: {str(e)}")
            raise e

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX files."""
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()

    def extract_text_from_txt(self, txt_path: str) -> str:
        """Extract text from TXT files."""
        with open(txt_path, "r", encoding="utf-8") as f:
            return f.read().strip()

    def extract_text_from_image_with_openai(self, image_path: str, model: str = "gpt-4o-mini") -> str:
        """Extract text from images using OpenAI Vision."""
        with open(image_path, "rb") as f:
            img_bytes = f.read()
        img_b64 = base64.b64encode(img_bytes).decode("utf-8")

        response = self.client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Extract all text clearly from this image."},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"},
                        },
                    ],
                }
            ],
        )

        return response.choices[0].message.content.strip()

    def extract_text(self, file_path: str, model: str = "gpt-4o-mini") -> str:
        """Auto-detect file type and extract text."""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            logger.info(f"Extracting text from {ext} file: {os.path.basename(file_path)}")

            if ext == ".pdf":
                return self.extract_text_from_pdf(file_path, model)
            elif ext == ".docx":
                return self.extract_text_from_docx(file_path)
            elif ext == ".txt":
                return self.extract_text_from_txt(file_path)
            elif ext in [".jpg", ".jpeg", ".png", ".jfif", ".bmp", ".tiff", ".tif", ".webp", ".heic"]:
                return self.extract_text_from_image_with_openai(file_path, model)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise e

    def process_text_with_prompt(self, text: str, prompt_template: str, model: str = "gpt-4o-mini") -> Dict[str, Any]:
        """Process extracted text using a custom prompt template."""
        # Replace {text} placeholder in the prompt template with actual text
        prompt = prompt_template.format(text=text)
        
        # For very long texts, we might need to chunk them
        max_text_length = 100000  # 100k characters limit for OpenAI
        if len(text) > max_text_length:
            logger.warning(f"Text is very long ({len(text)} chars), truncating to {max_text_length} chars")
            text = text[:max_text_length] + "\n\n[Text truncated due to length]"
            prompt = prompt_template.format(text=text)
        
        # Measure processing time
        start_time = time.time()
        
        response = self.client.chat.completions.create(
            model=model, 
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        
        end_time = time.time()
        elapsed = end_time - start_time
        content = response.choices[0].message.content.strip()

        # Strip markdown code blocks if present
        if content.startswith("```"):
            content = "\n".join(content.split("\n")[1:])
            if content.endswith("```"):
                content = "\n".join(content.split("\n")[:-1])

        try:
            parsed = json.loads(content)
        except Exception as e:
            logger.error(f"Error converting text to JSON: {e}")
            parsed = {"raw_output": content, "error": str(e)}

        # Log usage statistics
        usage = response.usage
        logger.info(f"Processing completed - Total tokens: {usage.total_tokens}, Time: {elapsed:.2f}s")

        return {
            "result": parsed,
            "usage": {
                "total_tokens": usage.total_tokens,
                "processing_time": elapsed
            }
        }

    def process_document(self, db: Session, user_id: str, prompt_id: str, file_path: str, original_filename: str) -> ProcessedDocument:
        """Process a document with a specific prompt and save results to database."""
        try:
            # Get the prompt from database
            prompt = db.query(DynamicPrompt).filter(
                DynamicPrompt.id == prompt_id,
                DynamicPrompt.user_id == user_id,
                DynamicPrompt.is_active == True
            ).first()
            
            if not prompt:
                raise ValueError("Prompt not found or not active")

            # Create processing record
            file_type = os.path.splitext(file_path)[1].lower()
            processed_doc = ProcessedDocument(
                user_id=user_id,
                prompt_id=prompt_id,
                original_filename=original_filename,
                file_path=file_path,
                file_type=file_type,
                processing_status="processing"
            )
            db.add(processed_doc)
            db.commit()
            db.refresh(processed_doc)

            # Extract text from document
            logger.info(f"Starting text extraction for {original_filename}")
            extracted_text = self.extract_text(file_path, prompt.gpt_model)
            
            if not extracted_text or len(extracted_text.strip()) < 10:
                raise ValueError("No meaningful text extracted from document")
            
            # Update with extracted text
            processed_doc.extracted_text = extracted_text
            db.commit()
            logger.info(f"Text extraction completed. Extracted {len(extracted_text)} characters")

            # Process text with custom prompt
            logger.info(f"Processing text with prompt: {prompt.name} using model: {prompt.gpt_model}")
            result = self.process_text_with_prompt(extracted_text, prompt.prompt_template, prompt.gpt_model)
            
            # Update with final result
            processed_doc.processed_result = json.dumps(result["result"])
            processed_doc.processing_status = "completed"
            db.commit()

            logger.info(f"Document processing completed successfully for {original_filename}")
            return processed_doc

        except Exception as e:
            logger.error(f"Error processing document {original_filename}: {str(e)}")
            
            # Update processing record with error
            if 'processed_doc' in locals():
                processed_doc.processing_status = "failed"
                processed_doc.error_message = str(e)
                db.commit()
            
            raise e

    def get_file_type(self, file_path: str) -> str:
        """Get file type from file path."""
        return os.path.splitext(file_path)[1].lower()
